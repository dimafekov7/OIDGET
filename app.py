import streamlit as st
import requests
import warnings
from requests.packages.urllib3.exceptions import InsecureRequestWarning
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io

warnings.simplefilter('ignore', InsecureRequestWarning)


@st.cache_data
def get_org_info(inn):
    base_url = "https://nsi.rosminzdrav.ru/api/data"
    params = {
        "identifier": "1.2.643.5.1.13.13.11.1461",
        # version убрали — берёт актуальную
        "query": inn,
        "page": "1",
        "size": "10",
        "queryCount": "true"
    }
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT )
// 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "application/json, text/plain, */*",
        "Referer": "https://nsi.rosminzdrav.ru/",
    }
    
    # Прокси в формате http://ip:port
    proxies = {
        "http": "http://79.174.12.190:80",
        "https": "http://79.174.12.190:80",
    }
    
    for attempt in range(4):  # 4 попытки
        try:
            response = requests.get(
                base_url,
                params=params,
                headers=headers,
                proxies=proxies,
                verify=False,
                timeout=60
            )
            if response.status_code == 200:
                data = response.json()
                active_orgs = [org for org in data.get("list", []) if not org.get("deleteDate")]
                if not active_orgs:
                    raise Exception("Нет активной организации по этому ИНН")
                
                active_orgs.sort(key=lambda x: x.get("modifyDate", ""), reverse=True)
                org = active_orgs[0]
                
                return org.get("oid"), org.get("nameFull")
                
        except Exception as e:
            if attempt == 3:
                raise Exception(f"Не удалось получить данные через прокси (попытка {attempt+1}): {str(e)}")
            continue  # пробуем ещё раз
    
    raise Exception("Не удалось подключиться к nsi.rosminzdrav.ru")    return org.get("oid"), org.get("nameFull")


def fill_document(template_bytes, oid, name):
    doc = Document(io.BytesIO(template_bytes))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text == '?':
                    new_text = oid.strip()
                elif text == '!':
                    new_text = name.strip()
                else:
                    continue

                for p in cell.paragraphs[:]:
                    cell._element.remove(p._element)

                p = cell.add_paragraph(new_text)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.left_indent = Pt(0)

                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


st.title("Заполнение заявок по ИНН")

st.write("Введите ИНН медицинской организации — и получите сразу два заполненных документа.")

inn = st.text_input("ИНН (12 цифр)", max_chars=12)

if st.button("Заполнить и скачать документы"):
    if not inn or not inn.isdigit() or len(inn) not in (10, 12):
        st.error("Введите корректный ИНН (10 или 12 цифр)")
    else:
        with st.spinner("Получаем данные из ФРМО и заполняем документы..."):
            try:
                oid, name = get_org_info(inn)
                st.success(f"Найдено: {name}\nOID: {oid}")

                # Загружаем шаблоны из файлов (они должны лежать рядом с app.py)
                with open("REMD.docx", "rb") as f:
                    remd_template = f.read()
                with open("IEMK.docx", "rb") as f:
                    iemk_template = f.read()

                filled_remd = fill_document(remd_template, oid, name)
                filled_iemk = fill_document(iemk_template, oid, name)

                st.download_button(
                    label="Скачать Заявку на регистрацию (REMD)",
                    data=filled_remd,
                    file_name=f"filled_reg_{inn}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.download_button(
                    label="Скачать Заявку на доступ к ИЭМК",
                    data=filled_iemk,
                    file_name=f"filled_access_{inn}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"Ошибка: {str(e)}")

st.markdown("---")
st.caption("Работает на Streamlit • Шаблоны: REMD.docx и IEMK.docx должны быть в той же папке")